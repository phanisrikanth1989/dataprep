from docx import Document

from agents.tools.extract_doc import _read_sections


def test_read_sections_associates_tables_with_headings(tmp_path):
    doc = Document()
    doc.add_heading("Inputs and Schema", level=1)
    t1 = doc.add_table(rows=1, cols=1)
    t1.rows[0].cells[0].text = "schema-table"
    doc.add_heading("Sample Input", level=1)
    doc.add_heading("ledger", level=2)
    t2 = doc.add_table(rows=1, cols=1)
    t2.rows[0].cells[0].text = "ledger-table"
    path = tmp_path / "d.docx"
    doc.save(str(path))

    sections = _read_sections(Document(str(path)))

    assert set(sections.keys()) == {"Inputs and Schema", "Sample Input"}
    assert sections["Inputs and Schema"][0][0] is None
    assert sections["Inputs and Schema"][0][1].rows[0].cells[0].text == "schema-table"
    assert sections["Sample Input"][0][0] == "ledger"
    assert sections["Sample Input"][0][1].rows[0].cells[0].text == "ledger-table"


from agents.tools.extract_doc import _parse_schema_table


def _table(doc, header, rows):
    t = doc.add_table(rows=1, cols=len(header))
    for i, h in enumerate(header):
        t.rows[0].cells[i].text = h
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = v
    return t


def test_parse_schema_table_groups_by_source():
    doc = Document()
    t = _table(
        doc,
        ["Source", "Column", "Type", "Nullable", "Key"],
        [
            ["ledger", "txn_id", "str", "false", "true"],
            ["ledger", "amt", "float", "false", "false"],
            ["statement", "ref_id", "str", "false", "true"],
        ],
    )
    schema = _parse_schema_table(t)
    assert list(schema.keys()) == ["ledger", "statement"]
    assert schema["ledger"][0].name == "txn_id"
    assert schema["ledger"][0].type == "str"
    assert schema["ledger"][0].nullable is False
    assert schema["ledger"][0].key is True
    assert schema["ledger"][1].key is False


from agents.tools.extract_doc import _parse_rules_table


def test_parse_rules_table():
    doc = Document()
    t = _table(
        doc,
        ["ID", "Kind", "Description"],
        [
            ["R1", "Match", "match ledger.txn_id to statement.ref_id"],
            ["R2", "Tolerance", "amounts equal within 0.01"],
            ["", "", ""],  # blank row ignored
        ],
    )
    rules = _parse_rules_table(t)
    assert len(rules) == 2
    assert rules[0] == {"id": "R1", "kind": "match", "description": "match ledger.txn_id to statement.ref_id"}
    assert rules[1]["kind"] == "tolerance"


from agents.tools.extract_doc import _parse_data_block


def _named_table(doc, header, rows):
    return _table(doc, header, rows)


def test_parse_data_block_sample_input():
    doc = Document()
    items = [
        ("ledger", _named_table(doc, ["txn_id", "amt"], [["T1", "100.00"], ["T2", "50.00"]])),
        ("statement", _named_table(doc, ["ref_id", "amt"], [["T1", "100.00"]])),
    ]
    data = _parse_data_block(items)
    assert data["ledger"] == [{"txn_id": "T1", "amt": "100.00"}, {"txn_id": "T2", "amt": "50.00"}]
    assert data["statement"] == [{"ref_id": "T1", "amt": "100.00"}]


def test_parse_data_block_expected_output_extracts_composite_key():
    doc = Document()
    items = [
        ("matched", _named_table(doc, ["txn_id*", "src*", "amt"], [["T1", "ledger", "100.00"]])),
    ]
    data, keys = _parse_data_block(items, extract_keys=True)
    assert keys["matched"] == ["txn_id", "src"]
    assert data["matched"] == [{"txn_id": "T1", "src": "ledger", "amt": "100.00"}]


from agents.tools.extract_doc import compute_derived_facts


def test_compute_derived_facts():
    sample = {
        "ledger": [
            {"txn_id": "T1", "amt": "100"},
            {"txn_id": "T2", "amt": ""},
            {"txn_id": "T2", "amt": "50"},
        ]
    }
    facts = compute_derived_facts(sample)["ledger"]
    assert facts["txn_id"]["unique"] is False          # T2 repeats
    assert facts["txn_id"]["max_group_size"] == 2
    assert facts["txn_id"]["n_distinct"] == 2
    assert facts["txn_id"]["null_rate"] == 0.0
    assert facts["amt"]["null_rate"] == round(1 / 3, 4)  # one empty cell
    assert facts["amt"]["unique"] is True                # 100, 50 distinct among non-null


from agents.tools.extract_doc import _check_conformance


def test_conformance_ok():
    sections = {b: [] for b in ("Inputs and Schema", "Transformation Rules", "Sample Input", "Expected Output")}
    report = _check_conformance(
        sections,
        sources_schema={"ledger": ["x"]},
        rules=[{"id": "R1"}],
        sample_input={"ledger": [{"a": "1"}]},
        expected_output={"matched": [{"a": "1"}]},
    )
    assert report.ok is True


def test_conformance_missing_block():
    sections = {"Inputs and Schema": [], "Transformation Rules": [], "Sample Input": []}
    report = _check_conformance(sections, {"ledger": ["x"]}, [{"id": "R1"}], {"ledger": [{"a": "1"}]}, {})
    assert report.ok is False
    assert "Expected Output" in report.missing_blocks


def test_conformance_empty_table_is_parse_error():
    sections = {b: [] for b in ("Inputs and Schema", "Transformation Rules", "Sample Input", "Expected Output")}
    report = _check_conformance(sections, {"ledger": ["x"]}, [{"id": "R1"}], {"ledger": []}, {"matched": [{"a": "1"}]})
    assert report.ok is False
    assert any("Sample Input" in e for e in report.parse_errors)


import pytest

from agents.tools.extract_doc import ConformanceError, extract_doc


def _build_recon_docx(path):
    doc = Document()
    doc.add_heading("Inputs and Schema", level=1)
    _table(doc, ["Source", "Column", "Type", "Nullable", "Key"], [
        ["ledger", "txn_id", "str", "false", "true"],
        ["ledger", "amt", "float", "false", "false"],
        ["statement", "ref_id", "str", "false", "true"],
        ["statement", "amt", "float", "false", "false"],
    ])
    doc.add_heading("Transformation Rules", level=1)
    _table(doc, ["ID", "Kind", "Description"], [
        ["R1", "match", "match ledger.txn_id to statement.ref_id"],
        ["R2", "tolerance", "amounts equal within 0.01"],
    ])
    doc.add_heading("Sample Input", level=1)
    doc.add_heading("ledger", level=2)
    _table(doc, ["txn_id", "amt"], [["T1", "100.00"], ["T2", "50.00"]])
    doc.add_heading("statement", level=2)
    _table(doc, ["ref_id", "amt"], [["T1", "100.00"]])
    doc.add_heading("Expected Output", level=1)
    doc.add_heading("matched", level=2)
    _table(doc, ["txn_id*", "amt"], [["T1", "100.00"]])
    doc.add_heading("breaks", level=2)
    _table(doc, ["txn_id*", "reason"], [["T2", "no_match"]])
    doc.save(str(path))


def test_extract_doc_end_to_end(tmp_path):
    path = tmp_path / "recon.docx"
    _build_recon_docx(path)

    result = extract_doc(str(path))

    assert result.conformance.ok is True
    assert set(result.sources_schema) == {"ledger", "statement"}
    assert result.rules[1]["kind"] == "tolerance"
    assert result.sample_input["ledger"][0] == {"txn_id": "T1", "amt": "100.00"}
    assert result.output_keys["matched"] == ["txn_id"]
    assert result.expected_output["breaks"][0] == {"txn_id": "T2", "reason": "no_match"}
    assert result.derived_facts["ledger"]["txn_id"]["unique"] is True


def test_extract_doc_raises_on_missing_block(tmp_path):
    doc = Document()
    doc.add_heading("Inputs and Schema", level=1)
    _table(doc, ["Source", "Column", "Type", "Nullable", "Key"], [["ledger", "txn_id", "str", "false", "true"]])
    path = tmp_path / "bad.docx"
    doc.save(str(path))

    with pytest.raises(ConformanceError) as exc:
        extract_doc(str(path))
    assert "Transformation Rules" in exc.value.report.missing_blocks


def test_extract_doc_raises_when_file_too_large(tmp_path, monkeypatch):
    path = tmp_path / "recon.docx"
    _build_recon_docx(path)
    monkeypatch.setattr("agents.tools.extract_doc.MAX_DOCX_BYTES", 10)

    with pytest.raises(ConformanceError) as exc:
        extract_doc(str(path))
    assert exc.value.report.ok is False
    assert any("file too large" in e for e in exc.value.report.parse_errors)


def _empty_table(doc, cols):
    return doc.add_table(rows=0, cols=cols)


def test_extract_doc_raises_with_parse_errors_on_degenerate_blocks(tmp_path):
    # All four blocks present, but every table is empty/image-only (no parseable
    # content), so the gate fails via parse_errors rather than missing_blocks.
    doc = Document()
    doc.add_heading("Inputs and Schema", level=1)
    _empty_table(doc, 5)
    doc.add_heading("Transformation Rules", level=1)
    _empty_table(doc, 3)
    doc.add_heading("Sample Input", level=1)
    _table(doc, ["txn_id", "amt"], [["T1", "1"]])  # orphan table: no Heading-2 -> skipped
    doc.add_heading("Expected Output", level=1)
    doc.add_heading("matched", level=2)
    _empty_table(doc, 2)
    path = tmp_path / "degenerate.docx"
    doc.save(str(path))

    with pytest.raises(ConformanceError) as exc:
        extract_doc(str(path))
    errors = exc.value.report.parse_errors
    assert exc.value.report.missing_blocks == []
    assert any("Inputs and Schema" in e for e in errors)
    assert any("Transformation Rules" in e for e in errors)
    assert any("Sample Input" in e for e in errors)
    assert any("Expected Output" in e for e in errors)


def test_extract_doc_raise_on_error_false_returns_nonok_result(tmp_path):
    doc = Document()
    doc.add_heading("Inputs and Schema", level=1)
    _table(doc, ["Source", "Column", "Type", "Nullable", "Key"], [["ledger", "txn_id", "str", "false", "true"]])
    path = tmp_path / "partial.docx"
    doc.save(str(path))

    result = extract_doc(str(path), raise_on_error=False)

    assert result.conformance.ok is False
    assert "Transformation Rules" in result.conformance.missing_blocks
    assert set(result.sources_schema) == {"ledger"}
    assert result.rules == []


def test_parse_schema_table_skips_blank_source_or_column():
    doc = Document()
    t = _table(
        doc,
        ["Source", "Column", "Type", "Nullable", "Key"],
        [
            ["ledger", "txn_id", "str", "false", "true"],
            ["", "orphan", "str", "true", "false"],   # blank Source -> skipped
            ["ledger", "", "str", "true", "false"],    # blank Column -> skipped
        ],
    )
    schema = _parse_schema_table(t)
    assert list(schema.keys()) == ["ledger"]
    assert len(schema["ledger"]) == 1


from types import SimpleNamespace

from agents.tools.extract_doc import _heading_level


def test_heading_level_edge_cases():
    assert _heading_level(SimpleNamespace(style=SimpleNamespace(name="Normal"))) is None       # non-heading
    assert _heading_level(SimpleNamespace(style=SimpleNamespace(name="Heading X"))) is None    # non-numeric suffix
    assert _heading_level(SimpleNamespace(style=None)) is None                                  # no style
    assert _heading_level(SimpleNamespace(style=SimpleNamespace(name="Heading 3"))) == 3        # valid heading


# ---------------------------------------------------------------------------
# Important #1 -- no silent row loss when a name/section repeats across tables.
# ---------------------------------------------------------------------------


def test_parse_data_block_accumulates_rows_for_repeated_name():
    # Two "ledger" tables (e.g. two Heading-2 "ledger" blocks under Sample Input)
    # must contribute ALL their rows, not just the last table's.
    doc = Document()
    items = [
        ("ledger", _named_table(doc, ["txn_id", "amt"], [["T1", "100.00"], ["T2", "50.00"]])),
        ("ledger", _named_table(doc, ["txn_id", "amt"], [["T3", "25.00"], ["T4", "10.00"]])),
    ]
    data = _parse_data_block(items)
    assert data["ledger"] == [
        {"txn_id": "T1", "amt": "100.00"},
        {"txn_id": "T2", "amt": "50.00"},
        {"txn_id": "T3", "amt": "25.00"},
        {"txn_id": "T4", "amt": "10.00"},
    ]


def test_parse_data_block_repeated_name_keeps_first_keys_and_all_rows():
    doc = Document()
    items = [
        ("matched", _named_table(doc, ["txn_id*", "amt"], [["T1", "100.00"]])),
        ("matched", _named_table(doc, ["txn_id", "amt"], [["T2", "50.00"]])),  # 2nd table: no key marker
    ]
    data, keys = _parse_data_block(items, extract_keys=True)
    assert keys["matched"] == ["txn_id"]  # first occurrence's key columns retained
    assert data["matched"] == [
        {"txn_id": "T1", "amt": "100.00"},
        {"txn_id": "T2", "amt": "50.00"},
    ]


def test_extract_doc_accumulates_repeated_sample_input_heading(tmp_path):
    # Two "ledger" Heading-2 blocks under "Sample Input" -> both tables' rows kept.
    doc = Document()
    doc.add_heading("Inputs and Schema", level=1)
    _table(doc, ["Source", "Column", "Type", "Nullable", "Key"], [["ledger", "txn_id", "str", "false", "true"]])
    doc.add_heading("Transformation Rules", level=1)
    _table(doc, ["ID", "Kind", "Description"], [["R1", "match", "m"]])
    doc.add_heading("Sample Input", level=1)
    doc.add_heading("ledger", level=2)
    _table(doc, ["txn_id", "amt"], [["T1", "100.00"], ["T2", "50.00"]])
    doc.add_heading("ledger", level=2)  # same source name, second table
    _table(doc, ["txn_id", "amt"], [["T3", "25.00"]])
    doc.add_heading("Expected Output", level=1)
    doc.add_heading("matched", level=2)
    _table(doc, ["txn_id*", "amt"], [["T1", "100.00"]])
    path = tmp_path / "dup.docx"
    doc.save(str(path))

    result = extract_doc(str(path))

    assert [r["txn_id"] for r in result.sample_input["ledger"]] == ["T1", "T2", "T3"]


def test_extract_doc_merges_schema_split_across_tables(tmp_path):
    # "Inputs and Schema" split across two tables -> all sources/columns present.
    doc = Document()
    doc.add_heading("Inputs and Schema", level=1)
    _table(doc, ["Source", "Column", "Type", "Nullable", "Key"], [
        ["ledger", "txn_id", "str", "false", "true"],
        ["statement", "ref_id", "str", "false", "true"],
    ])
    _table(doc, ["Source", "Column", "Type", "Nullable", "Key"], [
        ["ledger", "amt", "float", "false", "false"],
        ["statement", "amt", "float", "false", "false"],
    ])
    doc.add_heading("Transformation Rules", level=1)
    _table(doc, ["ID", "Kind", "Description"], [["R1", "match", "m"]])
    doc.add_heading("Sample Input", level=1)
    doc.add_heading("ledger", level=2)
    _table(doc, ["txn_id", "amt"], [["T1", "100.00"]])
    doc.add_heading("statement", level=2)
    _table(doc, ["ref_id", "amt"], [["T1", "100.00"]])
    doc.add_heading("Expected Output", level=1)
    doc.add_heading("matched", level=2)
    _table(doc, ["txn_id*", "amt"], [["T1", "100.00"]])
    path = tmp_path / "split.docx"
    doc.save(str(path))

    result = extract_doc(str(path))

    assert set(result.sources_schema) == {"ledger", "statement"}
    assert [c.name for c in result.sources_schema["ledger"]] == ["txn_id", "amt"]
    assert [c.name for c in result.sources_schema["statement"]] == ["ref_id", "amt"]


def test_extract_doc_concatenates_rules_across_tables(tmp_path):
    # "Transformation Rules" split across two tables -> rules concatenated.
    doc = Document()
    doc.add_heading("Inputs and Schema", level=1)
    _table(doc, ["Source", "Column", "Type", "Nullable", "Key"], [["ledger", "txn_id", "str", "false", "true"]])
    doc.add_heading("Transformation Rules", level=1)
    _table(doc, ["ID", "Kind", "Description"], [["R1", "match", "m1"]])
    _table(doc, ["ID", "Kind", "Description"], [["R2", "tolerance", "t1"]])
    doc.add_heading("Sample Input", level=1)
    doc.add_heading("ledger", level=2)
    _table(doc, ["txn_id"], [["T1"]])
    doc.add_heading("Expected Output", level=1)
    doc.add_heading("matched", level=2)
    _table(doc, ["txn_id*"], [["T1"]])
    path = tmp_path / "rules.docx"
    doc.save(str(path))

    result = extract_doc(str(path))

    assert [r["id"] for r in result.rules] == ["R1", "R2"]


# ---------------------------------------------------------------------------
# Important #3 -- open/parse failures are converted to ConformanceError.
# ---------------------------------------------------------------------------


def test_extract_doc_raises_conformance_on_non_docx_blob(tmp_path):
    path = tmp_path / "notdocx.docx"
    path.write_bytes(b"not a docx")
    with pytest.raises(ConformanceError) as exc:
        extract_doc(str(path))
    assert exc.value.report.ok is False
    assert any("could not open or parse docx" in e for e in exc.value.report.parse_errors)


def test_extract_doc_raises_conformance_on_missing_file(tmp_path):
    path = tmp_path / "does_not_exist.docx"
    with pytest.raises(ConformanceError) as exc:
        extract_doc(str(path))
    assert exc.value.report.ok is False
    assert any("could not open or parse docx" in e for e in exc.value.report.parse_errors)


def test_extract_doc_non_docx_returns_nonok_result_when_not_raising(tmp_path):
    path = tmp_path / "notdocx.docx"
    path.write_bytes(b"not a docx")
    result = extract_doc(str(path), raise_on_error=False)
    assert result.conformance.ok is False
    assert result.sources_schema == {}
    assert result.rules == []
    assert result.sample_input == {}
    assert result.expected_output == {}
