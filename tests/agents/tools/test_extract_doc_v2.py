from docx import Document

from agents.tools.extract_doc import (
    ExtractResult, _read_section_prose, extract_doc, to_dict,
)


def _table(doc, header, rows):
    t = doc.add_table(rows=1, cols=len(header))
    for i, h in enumerate(header):
        t.rows[0].cells[i].text = h
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = v
    return t


def test_extract_result_new_fields_default():
    r = ExtractResult(
        sources_schema={}, rules=[], sample_input={}, expected_output={},
        output_keys={}, derived_facts={}, conformance=None,
    )
    assert r.notes == "" and r.extra_sections == {} and r.tier == "build"
    d = to_dict(r)
    assert d["notes"] == "" and d["extra_sections"] == {} and d["tier"] == "build"


def test_read_section_prose_collects_paragraphs(tmp_path):
    doc = Document()
    doc.add_heading("Notes / Special Handling", level=1)
    doc.add_paragraph("Trim whitespace on all keys.")
    doc.add_paragraph("Amounts are in minor units.")
    path = tmp_path / "n.docx"
    doc.save(str(path))
    prose = _read_section_prose(Document(str(path)))
    assert prose["Notes / Special Handling"] == "Trim whitespace on all keys.\nAmounts are in minor units."


def _min_doc(path, with_notes=False):
    doc = Document()
    doc.add_heading("Inputs and Schema", level=1)
    _table(doc, ["Source", "Column", "Type", "Nullable", "Key"], [["src", "id", "str", "false", "true"]])
    doc.add_heading("Transformation Rules", level=1)
    _table(doc, ["ID", "Kind", "Description"], [["R1", "sort", "order by id"]])
    if with_notes:
        doc.add_heading("Notes / Special Handling", level=1)
        doc.add_paragraph("Keys are case-sensitive.")
    doc.save(str(path))


def test_extract_doc_captures_notes(tmp_path):
    path = tmp_path / "d.docx"
    _min_doc(path, with_notes=True)
    result = extract_doc(str(path))
    assert result.notes == "Keys are case-sensitive."


def test_extract_doc_no_notes_is_empty_string(tmp_path):
    path = tmp_path / "d.docx"
    _min_doc(path, with_notes=False)
    assert extract_doc(str(path)).notes == ""


def test_extra_sections_are_data_blind(tmp_path):
    doc = Document()
    doc.add_heading("Inputs and Schema", level=1)
    _table(doc, ["Source", "Column", "Type", "Nullable", "Key"], [["src", "id", "str", "false", "true"]])
    doc.add_heading("Transformation Rules", level=1)
    _table(doc, ["ID", "Kind", "Description"], [["R1", "sort", "order by id"]])
    doc.add_heading("Reference Codes", level=1)  # UNRECOGNIZED H1
    doc.add_paragraph("Map region codes to names.")
    _table(doc, ["code", "region"], [["EU", "Europe"], ["NA", "North America"]])
    path = tmp_path / "d.docx"
    doc.save(str(path))

    result = extract_doc(str(path))
    extra = result.extra_sections["Reference Codes"]
    assert extra["prose"] == "Map region codes to names."
    tbl = extra["tables"][0]
    assert tbl["columns"] == ["code", "region"]
    assert tbl["row_count"] == 2
    assert "region" in tbl["facts"] and "code" in tbl["facts"]
    assert "review" in tbl["flag"]
    # DATA-WALL: no raw cell value ('EU'/'Europe'/'NA') leaks into extra_sections.
    import json
    blob = json.dumps(result.extra_sections)
    assert "Europe" not in blob and "North America" not in blob and "EU" not in blob


def test_extra_sections_excludes_recognized_blocks(tmp_path):
    path = tmp_path / "d.docx"
    _min_doc(path, with_notes=True)
    result = extract_doc(str(path))
    # Notes / Special Handling is recognized (-> notes), never an extra section.
    assert result.extra_sections == {}


import json as _json

from agents.tools import extract_doc as _ed


def _full_doc(path, with_sample=True, with_expected_rows=True):
    doc = Document()
    doc.add_heading("Inputs and Schema", level=1)
    _table(doc, ["Source", "Column", "Type", "Nullable", "Key"], [["src", "id", "str", "false", "true"]])
    doc.add_heading("Transformation Rules", level=1)
    _table(doc, ["ID", "Kind", "Description"], [["R1", "sort", "order by id"]])
    if with_sample:
        doc.add_heading("Sample Input", level=1)
        doc.add_heading("src", level=2)
        _table(doc, ["id"], [["1"], ["2"]])
    doc.add_heading("Expected Output", level=1)
    doc.add_heading("out", level=2)
    rows = [["1"], ["2"]] if with_expected_rows else []
    _table(doc, ["id*"], rows)
    doc.save(str(path))


def test_tier_verified_when_sample_and_graded_expected(tmp_path):
    p = tmp_path / "v.docx"; _full_doc(p, with_sample=True, with_expected_rows=True)
    assert extract_doc(str(p)).tier == "verified"


def test_tier_smoke_when_sample_only(tmp_path):
    p = tmp_path / "s.docx"; _full_doc(p, with_sample=True, with_expected_rows=False)
    r = extract_doc(str(p))
    assert r.tier == "smoke"


def test_tier_build_when_no_sample(tmp_path):
    p = tmp_path / "b.docx"; _full_doc(p, with_sample=False, with_expected_rows=True)
    assert extract_doc(str(p)).tier == "build"


def test_cli_emits_tier(tmp_path):
    p = tmp_path / "v.docx"; _full_doc(p, with_sample=True, with_expected_rows=True)
    out = tmp_path / "e.json"
    rc = _ed.main([str(p), "--out", str(out)])
    assert rc == 0
    assert _json.loads(out.read_text())["tier"] == "verified"
