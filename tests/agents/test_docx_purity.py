"""Guardrail tests for the docx_purity pre-branch scanner."""
import base64
import io

from docx import Document

from agents.tools.docx_purity import scan_purity

# 1x1 transparent PNG (verified-valid; python-docx 1.2.0 embeds it and it yields
# a word/media/ entry, which is what the trip test exercises).
_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAAC0lEQVR42mNgAAIAAAUAAen63NgAAAAASUVORK5CYII=")


def test_template_docx_not_headingless():
    r = scan_purity("agents/examples/sample_etl_requirements.docx")
    assert r["has_headingless_content"] is False
    assert r["conformance_fail"] is False


def test_image_docx_trips(tmp_path):
    doc = Document()
    doc.add_paragraph("prose only, no template headings")
    doc.add_picture(io.BytesIO(_PNG))
    p = tmp_path / "img.docx"
    doc.save(p)
    r = scan_purity(str(p))
    assert r["has_images"] is True and r["tripped"] is True
