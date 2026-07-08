"""Generate a COMPLEX, non-template real-BRD for stress-testing the doc-normalizer front door.

Unlike the author-to-template docs, this deliberately looks like a real business
requirements document: a title heading, prose overview, an STTM (source-to-target
mapping) table with ARBITRARY headers, prose transformation rules, a free-prose
"special handling" section, and an EMBEDDED SCREENSHOT (a rendered PNG of the
expected-output sample). It also drops sibling CSVs next to the docx (the exact
input data). Run it, then point the etl-orchestrator at the .docx in real-BRD mode.

Usage:
    python -m agents.examples.gen_complex_real_brd [OUTPUT_DIR]

Everything is ASCII-only and uses neutral ETL vocabulary (join/lookup/validate/
filter/sort/derive) so it stays clean of the biased-framing naming gate.
"""
from __future__ import annotations

import csv
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont


# ------------------------------------------------------------------
# 1. Sibling input data (the EXACT sources -> rung-1 CSVs)
# ------------------------------------------------------------------
_TRADES = [
    ["trade_id", "account_id", "symbol", "quantity", "price", "status", "trade_date"],
    ["T001", "A100", "AAPL", "100", "150.25", "SETTLED", "2026-06-01"],
    ["T002", "A101", "MSFT", "50", "410.50", "SETTLED", "2026-06-01"],
    ["T003", "A100", "GOOG", "20", "2750.00", "PENDING", "2026-06-02"],
    ["T004", "A102", "AAPL", "200", "151.00", "SETTLED", "2026-06-02"],
    ["T005", "A999", "TSLA", "10", "690.10", "settled", "2026-06-03"],
]
_ACCOUNTS = [
    ["account_id", "account_name", "region"],
    ["A100", "Alpha Capital", "NA"],
    ["A100", "Alpha Capital Intl", "NA"],  # duplicate key -> deliberate fan-out hazard
    ["A101", "Beta Partners", "EMEA"],
    ["A102", "Gamma Funds", "APAC"],
    # A999 intentionally absent -> a no-match row (kept with nulls per the notes)
]
_PRICES = [
    ["symbol", "closing_price"],
    ["AAPL", "150.80"],
    ["MSFT", "411.00"],
    ["GOOG", "2748.50"],
    ["TSLA", "689.00"],
]

# The expected-output SAMPLE that gets rendered to a screenshot (image -> rung-3a).
# Its presence-as-an-image is the point: the answer key is a screenshot, so the
# validator caps the tier at "smoke" and never grades it (the safety cap).
_EXPECTED_SAMPLE = [
    ["trade_id", "account_name", "region", "symbol", "market_value", "closing_price"],
    ["T004", "Gamma Funds", "APAC", "AAPL", "30200.00", "150.80"],
    ["T002", "Beta Partners", "EMEA", "MSFT", "20525.00", "411.00"],
    ["T001", "Alpha Capital", "NA", "AAPL", "15025.00", "150.80"],
    ["T005", "(null)", "(null)", "TSLA", "6901.00", "689.00"],
]

# The STTM mapping table -- ARBITRARY headers, the kind a real BA writes.
_STTM_HEADERS = ["Source System", "Source Field(s)", "Target Field", "Transformation Logic"]
_STTM_ROWS = [
    ["Trades", "trade_id", "trade_id", "Pass through (record key)."],
    ["Trades", "quantity, price", "market_value", "Derive market_value = quantity * price."],
    ["Accounts", "account_name", "account_name", "Lookup on account_id (left join; keep unmatched)."],
    ["Accounts", "region", "region", "Lookup on account_id."],
    ["Prices", "closing_price", "closing_price", "Lookup on symbol."],
    ["Trades", "status", "-", "Filter: keep rows where status = SETTLED (case-insensitive)."],
    ["Trades", "trade_date", "trade_date", "Validate format YYYY-MM-DD."],
    ["Derived", "market_value", "-", "Sort output by market_value, descending."],
]


def _render_screenshot(path: Path, table: list[list[str]]) -> None:
    """Render a small table to a PNG so the doc has a real 'screenshot of sample data'."""
    font = ImageFont.load_default()
    cols = len(table[0])
    col_w, row_h, pad = 130, 26, 8
    width = cols * col_w + 2 * pad
    height = len(table) * row_h + 2 * pad
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    for r, row in enumerate(table):
        y = pad + r * row_h
        if r == 0:
            draw.rectangle([pad, y, width - pad, y + row_h], fill=(220, 228, 240))
        for c in range(cols + 1):  # vertical grid
            x = pad + c * col_w
            draw.line([x, pad, x, height - pad], fill=(180, 180, 180))
        draw.line([pad, y, width - pad, y], fill=(180, 180, 180))
        for c, cell in enumerate(row):
            draw.text((pad + c * col_w + 5, y + 6), str(cell)[:18], fill="black", font=font)
    draw.line([pad, height - pad, width - pad, height - pad], fill=(180, 180, 180))
    img.save(path)


def _write_csv(path: Path, rows: list[list[str]]) -> None:
    with open(path, "w", newline="", encoding="utf-8") as fh:
        csv.writer(fh).writerows(rows)


def generate(out_dir: str) -> Path:
    """Write the complex BRD .docx + its sibling CSVs into out_dir; return the .docx path."""
    out = Path(out_dir)
    out.mkdir(parents=True, exist_ok=True)

    # sibling exact-data CSVs (the exploder inventories these from the docx's own dir)
    _write_csv(out / "trades.csv", _TRADES)
    _write_csv(out / "accounts.csv", _ACCOUNTS)
    _write_csv(out / "prices.csv", _PRICES)

    # the expected-output screenshot (embedded image -> rung-3a)
    shot = out / "_expected_sample.png"
    _render_screenshot(shot, _EXPECTED_SAMPLE)

    doc = Document()
    doc.add_heading("Trade Position Build - Business Requirements", level=1)

    doc.add_heading("1. Overview", level=2)
    doc.add_paragraph(
        "This job builds an enriched trade position feed for the downstream reporting team. "
        "It reads a daily trades extract, adds account and price attributes by lookup, computes "
        "a market value, keeps only settled trades, and delivers a sorted output file. Sample input "
        "data is attached as CSV files alongside this document (trades.csv, accounts.csv, prices.csv)."
    )

    doc.add_heading("2. Source-to-Target Mapping", level=2)
    doc.add_paragraph("The mapping below defines each target field and how it is produced.")
    tbl = doc.add_table(rows=1, cols=len(_STTM_HEADERS))
    tbl.style = "Table Grid"
    for i, h in enumerate(_STTM_HEADERS):
        tbl.rows[0].cells[i].text = h
    for row in _STTM_ROWS:
        cells = tbl.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v

    doc.add_heading("3. Business Rules", level=2)
    doc.add_paragraph("R1. Join each trade to Accounts on account_id to add account_name and region.")
    doc.add_paragraph("R2. Join each trade to Prices on symbol to add closing_price.")
    doc.add_paragraph("R3. Derive market_value as quantity multiplied by price.")
    doc.add_paragraph("R4. Filter the output to trades whose status is SETTLED.")
    doc.add_paragraph("R5. Validate that trade_date is a valid date in YYYY-MM-DD format.")
    doc.add_paragraph("R6. Sort the final output by market_value in descending order.")

    doc.add_heading("4. Special Handling", level=2)
    doc.add_paragraph(
        "A trade whose account_id has no matching account (for example a brand-new account not yet "
        "in the Accounts file) must be KEPT in the output with account_name and region left null - do "
        "not drop it."
    )
    doc.add_paragraph(
        "The status comparison is case-insensitive: a value of 'settled' must be treated the same as "
        "'SETTLED'."
    )
    doc.add_paragraph(
        "Note that account_id is NOT guaranteed unique in the Accounts file; the same account_id can "
        "appear more than once. Flag this if it affects the row count of the output."
    )

    doc.add_heading("5. Expected Result (sample screenshot)", level=2)
    doc.add_paragraph(
        "The screenshot below shows the expected output for the attached sample input (settled trades "
        "only, enriched and sorted by market_value descending). It is a pasted image, not a table."
    )
    doc.add_picture(str(shot), width=Inches(6.0))

    docx_path = out / "complex_real_brd.docx"
    doc.save(str(docx_path))
    shot.unlink(missing_ok=True)  # the screenshot now lives inside the .docx; drop the loose PNG
    return docx_path


def main(argv=None) -> int:
    argv = argv if argv is not None else sys.argv[1:]
    out_dir = argv[0] if argv else "agents/work/complex_brd"
    docx_path = generate(out_dir)
    print("wrote " + str(docx_path))
    print("siblings: trades.csv, accounts.csv, prices.csv (keep them in the SAME folder)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
