"""Generate a slightly-more-complex conformant ETL-requirements .docx for testing.

Step-up from the basic sample: 3 sources, TWO column-adding lookup joins
(counterparty + currency), a FILTER (keep only SETTLED), and a sort. The sample
data exercises: a counterparty no-match, a currency no-match, and a row the
filter drops -- all deterministic, so the Expected Output is an exact oracle.

Regenerate with:  python -m agents.examples.gen_sample_etl_requirements_complex
"""
from pathlib import Path

from docx import Document


def add_table(doc, rows):
    """Render a real Word table (Table Grid); row 0 is the header. '' = null cell."""
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            t.rows[r].cells[c].text = val
    return t


doc = Document()
doc.add_heading("Multi-Lookup Settlement Enrichment - Requirements", 0)
doc.add_paragraph(
    "Enrich the settlements feed with counterparty and currency reference data, "
    "keep only settled transactions, and sort. All source rows that pass the "
    "filter are kept; unmatched lookups leave the added column empty."
)

# ---- 1. Inputs and Schema ----
doc.add_heading("Inputs and Schema", 1)
add_table(doc, [
    ["Source", "Column", "Type", "Nullable", "Key"],
    ["transactions", "txn_id",            "str", "false", "true"],
    ["transactions", "counterparty_code", "str", "true",  "false"],
    ["transactions", "currency_code",     "str", "true",  "false"],
    ["transactions", "amount",            "str", "true",  "false"],
    ["transactions", "status",            "str", "false", "false"],
    ["counterparties", "counterparty_code", "str", "false", "true"],
    ["counterparties", "counterparty_name", "str", "false", "false"],
    ["currencies", "currency_code", "str", "false", "true"],
    ["currencies", "currency_name", "str", "false", "false"],
])

# ---- 2. Transformation Rules ----
doc.add_heading("Transformation Rules", 1)
add_table(doc, [
    ["ID", "Kind", "Description"],
    ["R1", "join",
     "Add counterparty_name from counterparties, matched on counterparty_code. "
     "Left join - keep every transaction; an unmatched counterparty_code leaves counterparty_name empty."],
    ["R2", "join",
     "Add currency_name from currencies, matched on currency_code. "
     "Left join - keep every transaction; an unmatched currency_code leaves currency_name empty."],
    ["R3", "filter",
     "Keep only rows where status equals SETTLED; drop all other statuses."],
    ["R4", "sort",
     "Order the result ascending by txn_id."],
])

# ---- 3. Sample Input ----
doc.add_heading("Sample Input", 1)

doc.add_heading("transactions", 2)
add_table(doc, [
    ["txn_id", "counterparty_code", "currency_code", "amount", "status"],
    ["T001", "CP1",   "USD", "100", "SETTLED"],
    ["T002", "CP2",   "EUR", "200", "SETTLED"],
    ["T003", "CP999", "USD", "300", "SETTLED"],   # counterparty no-match
    ["T004", "CP1",   "JPY", "400", "PENDING"],   # filtered out (not SETTLED)
    ["T005", "CP3",   "GBP", "500", "SETTLED"],   # currency no-match
])

doc.add_heading("counterparties", 2)
add_table(doc, [
    ["counterparty_code", "counterparty_name"],
    ["CP1", "Acme"],
    ["CP2", "Globex"],
    ["CP3", "Initech"],
])

doc.add_heading("currencies", 2)
add_table(doc, [
    ["currency_code", "currency_name"],
    ["USD", "US Dollar"],
    ["EUR", "Euro"],
    ["JPY", "Japanese Yen"],   # note: GBP intentionally absent -> T005 currency no-match
])

# ---- 4. Expected Output (deterministic oracle) ----
# After R1 (+counterparty_name), R2 (+currency_name), R3 (drop PENDING T004), R4 (sort txn_id):
doc.add_heading("Expected Output", 1)
doc.add_heading("enriched", 2)
add_table(doc, [
    ["txn_id*", "counterparty_code", "currency_code", "amount", "status", "counterparty_name", "currency_name"],
    ["T001", "CP1",   "USD", "100", "SETTLED", "Acme",    "US Dollar"],
    ["T002", "CP2",   "EUR", "200", "SETTLED", "Globex",  "Euro"],
    ["T003", "CP999", "USD", "300", "SETTLED", "",        "US Dollar"],   # cp no-match -> empty
    ["T005", "CP3",   "GBP", "500", "SETTLED", "Initech", ""],            # ccy no-match -> empty
    # T004 dropped by the filter (PENDING)
])

# ---- 5. Notes / Special Handling ----
doc.add_heading("Notes / Special Handling", 1)
doc.add_paragraph(
    "Keep every SETTLED transaction, even when a lookup does not match - an unmatched "
    "counterparty_code or currency_code leaves that added column empty rather than dropping the row. "
    "Non-SETTLED rows (e.g. PENDING) are removed entirely. counterparty_code and currency_code are "
    "case-sensitive."
)

out = Path("agents/examples/sample_etl_requirements_complex.docx")
out.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(out))
print("WROTE", out)
