"""Deterministically generate the neutral ETL example requirements .docx.

Run: python -m agents.examples.gen_sample_etl_requirements
Regenerate rather than hand-editing the tables when the schema changes.
"""
from pathlib import Path

from docx import Document

_OUT = Path("agents/examples/sample_etl_requirements.docx")


def _table(doc, header, rows):
    t = doc.add_table(rows=1, cols=len(header))
    for i, h in enumerate(header):
        t.rows[0].cells[i].text = h
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = v
    return t


def build():
    doc = Document()
    doc.add_heading("Inputs and Schema", level=1)
    _table(doc, ["Source", "Column", "Type", "Nullable", "Key"], [
        ["transactions", "txn_id", "str", "false", "true"],
        ["transactions", "counterparty_code", "str", "true", "false"],
        ["transactions", "amount", "str", "true", "false"],
        ["counterparties", "counterparty_code", "str", "false", "true"],
        ["counterparties", "counterparty_name", "str", "true", "false"],
    ])
    doc.add_heading("Transformation Rules", level=1)
    _table(doc, ["ID", "Kind", "Description"], [
        ["R1", "join", "add counterparty_name from counterparties on counterparty_code (left join, keep all txns)"],
        ["R2", "sort", "order the result by txn_id ascending"],
    ])
    doc.add_heading("Sample Input", level=1)
    doc.add_heading("transactions", level=2)
    _table(doc, ["txn_id", "counterparty_code", "amount"], [
        ["T001", "CP1", "100"], ["T002", "CP2", "200"], ["T003", "CP999", "300"],
    ])
    doc.add_heading("counterparties", level=2)
    _table(doc, ["counterparty_code", "counterparty_name"], [
        ["CP1", "Acme"], ["CP2", "Globex"],
    ])
    doc.add_heading("Expected Output", level=1)
    doc.add_heading("result", level=2)
    _table(doc, ["txn_id*", "counterparty_code", "amount", "counterparty_name"], [
        ["T001", "CP1", "100", "Acme"],
        ["T002", "CP2", "200", "Globex"],
        ["T003", "CP999", "300", ""],   # no counterparty match -> empty name, row kept
    ])
    doc.add_heading("Notes / Special Handling", level=1)
    doc.add_paragraph("Keep every transaction row; an unmatched counterparty_code leaves counterparty_name empty.")
    doc.add_paragraph("counterparty_code is case-sensitive.")
    _OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(_OUT))


if __name__ == "__main__":
    build()
