"""Deterministically generate a MESSY real-BRD .docx (+ sibling trades.csv) for the front-door E2E.

Unlike the template docs, this has NO template Heading-1 blocks: prose transform
rules, one STTM-style mapping table with arbitrary headers, and one embedded
screenshot-image of a small expected-result table. A sibling trades.csv is emitted
in the same output directory (the exploder's rung-1 CSV source). ASCII only.

Run: python -m agents.examples.gen_sample_real_brd --out-dir /tmp/real_brd_e2e
"""
from __future__ import annotations

import argparse
import base64
import io
from pathlib import Path

from docx import Document

# Corrected 1x1 transparent PNG (the SAME verified-valid constant python-docx can
# embed as tests/agents/test_docx_purity.py::_PNG; the plan's own base64 was corrupt).
_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAAC0lEQVR42mNgAAIAAAUAAen63NgAAAAASUVORK5CYII=")

# The sibling CSV: header + rows the exploder copies in as a rung-1 sample source.
_TRADES_CSV = (
    "trade_id,symbol,quantity,price,currency\n"
    "T001,aapl,10,150.00,USD\n"
    "T002,msft,5,300.00,USD\n"
    "T003,goog,2,2500.00,USD\n"
)


def _table(doc, header, rows):
    """Add a docx table: one header row then one row per data tuple."""
    t = doc.add_table(rows=1, cols=len(header))
    for i, h in enumerate(header):
        t.rows[0].cells[i].text = h
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = v
    return t


def build(out_dir) -> Path:
    """Write sample_real_brd.docx + sibling trades.csv into out_dir; return the docx path."""
    out = Path(out_dir)
    out.mkdir(parents=True, exist_ok=True)

    doc = Document()
    # ---- para:0 range: title + intro + prose rules + mapping lead-in (NO template blocks) ----
    doc.add_heading("FX Trade Amount Build - Business Requirements", level=1)
    doc.add_paragraph(
        "This job reads the daily trades file and builds a per-trade amount for downstream use.")
    doc.add_paragraph(
        "Rule R1: derive amount as quantity multiplied by price for every trade row.")
    doc.add_paragraph(
        "Rule R2: sort the result by trade_id in ascending order before hand-off.")
    doc.add_heading("Field Mapping", level=2)
    doc.add_paragraph(
        "The table below maps each source field to its target field and the transform to apply.")
    # ---- table:0 : STTM-style mapping table with arbitrary (non-template) headers ----
    _table(doc, ["Source Field", "Target Field", "Transformation Logic"], [
        ["trade_id", "trade_id", "direct copy; primary key"],
        ["symbol", "symbol", "uppercase the ticker"],
        ["quantity, price", "amount", "quantity multiplied by price"],
        ["currency", "(not carried)", "drop from the output"],
    ])
    # ---- para:1 range: expected-result lead-in + the screenshot image + notes ----
    doc.add_heading("Expected Result (screenshot)", level=2)
    doc.add_paragraph("A sample of the expected output is shown in the screenshot below.")
    doc.add_picture(io.BytesIO(_PNG))
    doc.add_paragraph(
        "Notes: keep every trade row; amount uses two decimal places; symbol is uppercased.")

    docx_path = out / "sample_real_brd.docx"
    doc.save(str(docx_path))
    (out / "trades.csv").write_text(_TRADES_CSV, encoding="utf-8")
    return docx_path


def main(argv=None) -> int:
    """CLI: generate the messy real-BRD docx + sibling trades.csv into --out-dir."""
    parser = argparse.ArgumentParser(description="Generate a messy real-BRD .docx + sibling trades.csv.")
    parser.add_argument("--out-dir", required=True, help="directory to write sample_real_brd.docx + trades.csv")
    args = parser.parse_args(argv)
    docx_path = build(args.out_dir)
    print(f"[gen_sample_real_brd] wrote {docx_path} + sibling trades.csv")
    return 0


if __name__ == "__main__":
    import sys

    sys.exit(main())
